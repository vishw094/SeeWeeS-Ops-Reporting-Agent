"""
Generates Technical_Business_Report.docx and converts it to PDF.
Run from project root:  python generate_report_docx.py
"""
from __future__ import annotations

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
NAVY        = RGBColor(0x00, 0x33, 0x66)
BLUE        = RGBColor(0x00, 0x55, 0x99)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE  = RGBColor(0xE8, 0xF2, 0xFF)
DARK_BG     = RGBColor(0x1E, 0x1E, 0x1E)
GREEN_PASS  = RGBColor(0x4E, 0xC9, 0x94)
GREY_TXT    = RGBColor(0x88, 0x88, 0x88)
BLUE_PROMPT = RGBColor(0x9C, 0xDC, 0xFE)
AMBER       = RGBColor(0x8B, 0x5E, 0x00)
GREEN_DARK  = RGBColor(0x1A, 0x6E, 0x2E)
RED_FLAG    = RGBColor(0xCC, 0x22, 0x00)
GREEN_OK    = RGBColor(0x22, 0x8B, 0x22)
LIGHT_GREY  = RGBColor(0xF5, 0xF5, 0xF5)


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------
def _set_cell_bg(cell, hex_color: str):
    """Fill a table cell with a solid background colour (hex without #)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.upper())
    tcPr.append(shd)


def _set_para_bg(para, hex_color: str):
    """Fill a paragraph's background (used inside table cells)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.upper())
    pPr.append(shd)


def _page_break(doc: Document):
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(docx_break_type("page"))


def docx_break_type(kind: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE


def add_page_break(doc: Document):
    from docx.enum.text import WD_BREAK
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(WD_BREAK.PAGE)


def set_col_width(table, col_idx: int, width: Inches):
    for row in table.rows:
        row.cells[col_idx].width = width


# ---------------------------------------------------------------------------
# Style shortcuts
# ---------------------------------------------------------------------------
def heading1(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "003366")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE
    return p


def heading3(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    return p


def body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.first_line_indent = Pt(0)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def bullet(doc: Document, text: str, bold_prefix: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        p.add_run(text).font.size = Pt(10.5)
    else:
        p.add_run(text).font.size = Pt(10.5)
    return p


def callout(doc: Document, text: str, color_hex="E8F2FF", border_hex="005599"):
    """Simulated callout using a 1-column borderless table with shading."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, color_hex)
    cell.width = Inches(6.0)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    # left border only
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "18")
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), border_hex)
    tcBdr.append(left)
    tcPr.append(tcBdr)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# Header table (dark blue)
# ---------------------------------------------------------------------------
def make_header_table(doc: Document, cols: list[str]):
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(cols):
        cell = tbl.rows[0].cells[i]
        _set_cell_bg(cell, "003366")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
    return tbl


def add_data_row(tbl, values: list[str], center_cols: list[int] = None, flag_col: int = None, ok_col: int = None, shade_even: bool = False, row_idx: int = 0):
    row = tbl.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        if shade_even and row_idx % 2 == 1:
            _set_cell_bg(cell, "F5F8FF")
        p = cell.paragraphs[0]
        if center_cols and i in center_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.font.size = Pt(9)
        if flag_col is not None and i == flag_col and val not in ("0", "No", "None", ""):
            run.font.color.rgb = RED_FLAG
            run.bold = True
        if ok_col is not None and i == ok_col and val in ("0", "No", "PASSED", "pass", "None"):
            run.font.color.rgb = GREEN_OK


# ---------------------------------------------------------------------------
# Terminal block
# ---------------------------------------------------------------------------
def add_terminal(doc: Document):
    """Simulate a VS Code dark terminal with coloured pytest output."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, "1E1E1E")

    def mono_line(cell, parts: list[tuple[str, RGBColor]], space_after=0):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)
        for text, color in parts:
            r = p.add_run(text)
            r.font.name = "Courier New"
            r.font.size = Pt(8)
            r.font.color.rgb = color
        return p

    D  = RGBColor(0xD4, 0xD4, 0xD4)   # default text
    G  = GREEN_PASS                     # PASSED
    GR = GREY_TXT                       # percentage
    B  = BLUE_PROMPT                    # prompt
    GN = RGBColor(0x6A, 0x99, 0x55)    # green comment / divider

    # Clear the auto paragraph
    cell.paragraphs[0].clear()

    mono_line(cell, [(r"PS D:\SeeWeeS Ops Reporting Agent> ", B), ("python", RGBColor(0x56,0x9C,0xD6)), (" -m pytest tests/ -v", D)])
    mono_line(cell, [("=" * 62, GN)])
    mono_line(cell, [("platform win32 -- Python 3.11.4, pytest-9.0.3, pluggy-1.6.0", D)])
    mono_line(cell, [("rootdir: D:\\SeeWeeS Ops Reporting Agent", D)])
    mono_line(cell, [("plugins: anyio-4.12.1, langsmith-0.8.4", D)])
    mono_line(cell, [("collected ", D), ("29", RGBColor(0xDC,0xDC,0xAA)), (" items", D)], space_after=3)

    tests = [
        ("tests/test_audit_logic.py::test_expected_buffer_mapping",                          "[  3%]"),
        ("tests/test_audit_logic.py::test_audit_passes_when_buffer_matches_and_rules_cited", "[  6%]"),
        ("tests/test_audit_logic.py::test_audit_flags_wrong_buffer",                         "[ 10%]"),
        ("tests/test_audit_logic.py::test_audit_requires_escalation_at_score_3",             "[ 13%]"),
        ("tests/test_audit_logic.py::test_audit_requires_cited_rules",                       "[ 17%]"),
        ("tests/test_audit_logic.py::test_llm_failure_keeps_failing_even_if_deterministic_passes", "[ 20%]"),
        ("tests/test_dq_tools.py::test_exact_match_is_valid",                                "[ 24%]"),
        ("tests/test_dq_tools.py::test_alias_match_resolves_when_item_id_agrees",            "[ 27%]"),
        ("tests/test_dq_tools.py::test_legacy_id_map",                                       "[ 31%]"),
        ("tests/test_dq_tools.py::test_name_alias_takes_precedence_over_legacy_id",          "[ 34%]"),
        ("tests/test_dq_tools.py::test_dq01_missing_uid_excluded",                           "[ 37%]"),
        ("tests/test_dq_tools.py::test_dq02_invalid_item_id_excluded",                       "[ 41%]"),
        ("tests/test_dq_tools.py::test_dq04_duplicate_uid_flagged",                          "[ 44%]"),
        ("tests/test_dq_tools.py::test_sla_tier_backfilled",                                 "[ 48%]"),
        ("tests/test_dq_tools.py::test_assign_tier_defaults_to_two",                         "[ 51%]"),
        ("tests/test_dq_tools.py::test_missing_required_columns_raises",                     "[ 55%]"),
        ("tests/test_resource_tools.py::test_load_resource_availability_parses_csv",         "[ 58%]"),
        ("tests/test_resource_tools.py::test_compute_demand_trucks_match_capacity_model",    "[ 62%]"),
        ("tests/test_resource_tools.py::test_compute_demand_cold_chain_uses_temp_trucks",    "[ 65%]"),
        ("tests/test_resource_tools.py::test_compute_demand_excludes_invalid_rows",          "[ 68%]"),
        ("tests/test_resource_tools.py::test_no_shortage_zero_penalty",                      "[ 72%]"),
        ("tests/test_resource_tools.py::test_temp_truck_shortage_incurs_cold_chain_penalty", "[ 75%]"),
        ("tests/test_resource_tools.py::test_tier1_penalised_more_than_tier2",               "[ 79%]"),
        ("tests/test_resource_tools.py::test_allocation_bottleneck_identified",              "[ 82%]"),
        ("tests/test_smoke.py::test_smoke",                                                  "[ 86%]"),
        ("tests/test_trend_tools.py::test_required_trucks_capacity_model",                   "[ 89%]"),
        ("tests/test_trend_tools.py::test_corridor_kpis_planning_window",                    "[ 93%]"),
        ("tests/test_trend_tools.py::test_pop_trend_available",                              "[ 96%]"),
        ("tests/test_trend_tools.py::test_deep_dive_tables_shapes",                          "[100%]"),
    ]

    for name, pct in tests:
        mono_line(cell, [(name + " ", D), ("PASSED", G), ("  " + pct, GR)])

    mono_line(cell, [("", D)], space_after=2)
    mono_line(cell, [("=" * 30 + " ", GN), ("29 passed", G), (" in 40.88s ", D), ("=" * 30, GN)])
    mono_line(cell, [(r"PS D:\SeeWeeS Ops Reporting Agent> ", B)])


# ---------------------------------------------------------------------------
# Architecture diagram box
# ---------------------------------------------------------------------------
def arch_box(doc: Document, lines: list[str]):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, "F9F9FF")
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "C0C8E0")
        tcBdr.append(el)
    tcPr.append(tcBdr)

    cell.paragraphs[0].clear()
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x44)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# Main document builder
# ---------------------------------------------------------------------------
def build(out_path: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default paragraph font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # ---------------------------------------------------------------- COVER
    cover_tbl = doc.add_table(rows=1, cols=1)
    cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cc = cover_tbl.rows[0].cells[0]
    _set_cell_bg(cc, "003366")
    cc.width = Inches(6.5)

    for text, size, bold, italic in [
        ("UCLA MSBA AI AGENTS PROJECT CHALLENGE 2026", 9, False, False),
        ("", 6, False, False),
        ("SeeWeeS Specialty Medicine Dispatch Network", 11, False, True),
        ("", 6, False, False),
        ("Multi-Agent Ops Reporting System", 22, True, False),
        ("", 4, False, False),
        ("Technical and Business Documentation", 13, False, True),
        ("", 10, False, False),
        ("Features: Self-Correction Audit Loop  |  Deep-Dive Trend Analysis  |  Multi-Region Resource Planning", 9, False, False),
        ("", 4, False, False),
        ("Platform: LangGraph 0.2  |  LangChain  |  GPT-4.1-mini  |  ChromaDB  |  Python 3.11", 9, False, False),
        ("", 4, False, False),
        ("May 2026", 10, False, False),
    ]:
        p = cc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(text)
        r.font.color.rgb = WHITE
        r.font.size  = Pt(size)
        r.bold       = bold
        r.italic     = italic

    add_page_break(doc)

    # ---------------------------------------------------------------- TOC (manual)
    heading1(doc, "Table of Contents")
    toc_entries = [
        ("1.  Executive Summary", ""),
        ("    Stakeholder and Business Context", ""),
        ("    Operational Pain Point", ""),
        ("    Solution at a Glance", ""),
        ("2.  Key Assumptions", ""),
        ("    Logistics Constraints", ""),
        ("    Data Availability", ""),
        ("    Business Rules", ""),
        ("3.  Technical Methodology", ""),
        ("    3.1  Architectural Enhancements", ""),
        ("    3.2  Agent Design and Specialized Prompts", ""),
        ("    3.3  Data Quality and Reconciliation Engine", ""),
        ("    3.4  Trend Analysis and KPI Computation", ""),
        ("    3.5  Resource Allocation", ""),
        ("4.  Results and Validation", ""),
        ("    4.1  Actual Pipeline Output", ""),
        ("    4.2  Data Quality Results", ""),
        ("    4.3  Item Spike Detection", ""),
        ("    4.4  Audit Loop Validation", ""),
        ("    4.5  Performance Benchmarks", ""),
        ("5.  Limitations and Next Steps", ""),
    ]
    for entry, _ in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(entry)
        r.font.size = Pt(10)
        if not entry.startswith("    "):
            r.bold = True
            r.font.color.rgb = NAVY
    add_page_break(doc)

    # ============================================================ SECTION 1
    heading1(doc, "1.  Executive Summary")

    heading2(doc, "Stakeholder and Business Context")
    body(doc, "The primary stakeholder for this system is the dispatch operations manager at SeeWeeS, a specialty medicine logistics company operating out of New Jersey distribution centers. The operations team is responsible for planning and executing time-critical medicine shipments to hospitals across two defined delivery corridors: the NJ to Boston corridor (C1_I95_NJ_BOS, Tier 1) and the NJ to Philadelphia corridor (C2_NJ_PHL, Tier 2). Both corridors carry life-critical medicines including antivirals, monoclonal antibodies, cold-chain biologics, and controlled substances, each with strict SLA requirements of 6 to 12 hours maximum transit time.")
    body(doc, "The operational decisions made each morning cover a 48-hour planning horizon (Day 0 and Day 1) and involve four interconnected problems: whether the incoming shipment data is accurate and complete, how many trucks and drivers are needed per corridor, what weather-related risks exist along each route, and whether the resulting dispatch plan is compliant with SeeWeeS policy before it leaves the operations manager's desk.")

    heading2(doc, "Operational Pain Point")
    body(doc, "Before this project, the reporting system was a generic, single-corridor linear pipeline. It ran an anomaly detection algorithm (IsolationForest) on whatever numeric columns happened to be present in the CSV, fetched weather for one hardcoded location, and passed a vague text summary to an LLM that produced an unverifiable plan. There was no mechanism to catch or fix planning errors before the report was finalized. Critically:")
    bullet(doc, "Incoming shipment rows with missing or mismatched item identifiers were silently included in dispatch totals, inflating truck and driver requirements.")
    bullet(doc, "There was no cross-corridor comparison. The NJ to Philadelphia route was not accounted for at all.")
    bullet(doc, "A planner that recommended the wrong travel buffer or ignored escalation requirements had no way to be corrected before the report reached leadership.")
    bullet(doc, "Every pipeline run re-extracted context from the policy document even when the document had not changed, wasting 15 to 20 seconds on every execution.")

    heading2(doc, "Solution at a Glance")
    body(doc, "This project delivers three integrated enhancements that address each pain point directly. Feature 1 adds a self-correcting audit loop that rejects non-compliant dispatch plans and forces the planner to revise before the report is generated, with human escalation when the loop cannot converge automatically. Feature 3 replaces the generic anomaly detection with a deterministic item master reconciliation engine and corridor-aware trend analysis grounded in the actual SeeWeeS Dispatch Playbook rules. Feature 5 extends the weather module to evaluate all nine corridor waypoints in parallel and adds a penalty-minimising resource allocator that computes the optimal truck and driver assignment across both corridors and both planning days.")

    # Stat summary table
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    stat_tbl = doc.add_table(rows=2, cols=5)
    stat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    stat_headers = ["Features Implemented", "Unit Tests Passing", "DQ Pass Rate", "Audit Attempts to Pass", "Cold Run Duration"]
    stat_values  = ["3", "29", "96.1%", "1", "~63 seconds"]
    stat_colors  = ["003366", "1A6E2E", "003366", "1A6E2E", "003366"]
    for i, (h, v, c) in enumerate(zip(stat_headers, stat_values, stat_colors)):
        hcell = stat_tbl.rows[0].cells[i]
        vcell = stat_tbl.rows[1].cells[i]
        _set_cell_bg(hcell, c)
        _set_cell_bg(vcell, c)
        hp = hcell.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(v)
        hr.font.size = Pt(18)
        hr.bold = True
        hr.font.color.rgb = WHITE
        vp = vcell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.add_run(h)
        vr.font.size = Pt(8)
        vr.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_page_break(doc)

    # ============================================================ SECTION 2
    heading1(doc, "2.  Key Assumptions")

    heading2(doc, "Logistics Constraints")
    bullet(doc, "Each standard truck holds exactly 10 volume units, where each unique shipment ID counts as one unit. A 10 percent packing inefficiency buffer is applied, giving a working formula of trucks required equals the ceiling of total units multiplied by 1.10 divided by 10. This comes directly from Playbook section 7.2.", "Truck capacity:  ")
    bullet(doc, "Items requiring Cold (2-8 degrees C) or Strict Cold Chain (minus 20 degrees C) must be transported exclusively in temperature-controlled refrigerator trucks. Standard trucks cannot be used for these items under any circumstance. This is a hard constraint, not a soft preference.", "Cold-chain separation:  ")
    bullet(doc, "One driver is required per truck, so total drivers required equals standard trucks plus temperature-controlled trucks.", "Driver assignment:  ")
    bullet(doc, "The resource pool (6 drivers, 4 standard trucks, 2 refrigerator trucks per day) is fixed and shared across both corridors simultaneously. There is no overflow buffer or on-call pool assumed.", "Resource pool:  ")
    bullet(doc, "The system plans for Day 0 and Day 1 only. Decisions for Day 2 and beyond are out of scope for this pipeline run.", "48-hour planning horizon:  ")

    heading2(doc, "Data Availability")
    bullet(doc, "The incoming CSV is assumed to be provided daily and covers 14 days of data (12 history days plus Day 0 and Day 1 for the planning window). The schema is fixed as defined in Playbook section 10.", "Shipment CSV:  ")
    bullet(doc, "Open-Meteo is used as the weather data source because it provides freely available daily aggregate forecasts (precipitation, wind gusts, minimum temperature) without requiring an API key. The two-day forecast window is used to cover Day 0 and Day 1 risk simultaneously.", "Weather API:  ")
    bullet(doc, "The canonical item master, alias table, and legacy ID mapping are treated as stable reference data and are hard-coded as module constants. This was a deliberate choice: an externally-managed CSV would introduce a dependency that could go stale between runs, whereas hard-coded constants are always in sync with the playbook version the code was built against.", "Item master:  ")
    bullet(doc, "The SeeWeeS Specialty Dispatch Playbook (markdown format) is the single source of truth for all thresholds, SLA definitions, and business rules. The RAG system indexes this document at startup and rebuilds the index only when the file changes.", "Policy document:  ")

    heading2(doc, "Business Rules")
    bullet(doc, "Tier 1 (life-critical, 6-hour maximum transit) covers antivirals, hormones, monoclonal antibodies, emergency drugs, anticoagulants, and clinical trial drugs. Tier 2 (standard specialty, 12-hour maximum) covers opioid analgesics and bronchodilators. Any medicine type not in the lookup defaults to Tier 2.", "SLA tiers are fixed by medicine type:  ")
    bullet(doc, "When resources are insufficient to serve all demand, Tier 1 cold-chain demand is served first (180 penalty points per unmet unit), then Tier 1 room-temperature (100 points), then Tier 2 cold-chain (120 points), then Tier 2 room-temperature (40 points). This ordering directly reflects the Playbook section 13.2 penalty model.", "Allocation priority:  ")
    bullet(doc, "Weather risk scores map to mandatory travel time buffers: score 0 means no buffer, score 1 means plus 10 percent, score 2 means plus 25 percent, and score 3 means plus 40 percent with mandatory escalation to a human manager.", "Travel buffer mapping:  ")
    bullet(doc, "Any row excluded by a data quality rule (DQ-01 through DQ-04) is completely removed from dispatch calculations. It is never counted in truck volumes, driver requirements, or SLA metrics. This is not configurable.", "DQ exclusions are hard:  ")
    add_page_break(doc)

    # ============================================================ SECTION 3
    heading1(doc, "3.  Technical Methodology")

    heading2(doc, "3.1  Architectural Enhancements")
    body(doc, "The original system used a simple linear LangGraph pipeline: context extraction, CSV analysis, weather fetch, planner, then report. The enhanced architecture introduces three structural changes that together produce a fundamentally more capable system.")

    heading3(doc, "Cyclic Audit Loop (Feature 1)")
    body(doc, "A conditional edge is added after the planner node. The AuditAgent inspects the dispatch plan against six Playbook rules and returns a structured JSON verdict. If the plan fails, the violations are injected back into the planner prompt as explicit correction requirements and the planner retries. This cycle continues for up to three attempts. After three failures, the routing logic checks the weather risk score: if it is at maximum (3), the pipeline escalates to a human checkpoint node that prints a structured alert and waits for manager approval via stdin.")
    arch_box(doc, [
        "Feature 1 Graph (cyclic):",
        "",
        "  node_planner  <---------- retry (violations fed back as prompt input)",
        "       |",
        "       v",
        "  node_audit",
        "       |",
        "       |-- passed = True              --> node_report",
        "       |",
        "       |-- retries < 3, failed        --> node_planner  (with violations list)",
        "       |",
        "       |-- retries >= 3, score = 3   --> node_human_checkpoint",
        "       |                                        |",
        "       |                                        +-- approved --> node_report",
        "       |",
        "       +-- retries >= 3, score < 3   --> node_report (with audit warnings)",
    ])

    heading3(doc, "Parallel Fan-Out for DQ and Weather (Features 3 and 5)")
    body(doc, "The DQ reconciliation node and the multi-corridor weather node are independent: the first reads only from the shipment CSV, and the second calls the Open-Meteo API for nine waypoint locations. In LangGraph 0.2, parallel branches are declared simply by adding two edges from the same source node. LangGraph executes them in separate threads and performs a deep merge of their returned state dictionaries before the fan-in node (resource allocator) can run.")
    arch_box(doc, [
        "Final Architecture (parallel + cyclic):",
        "",
        "  node_pdf_context",
        "        |",
        "        +---------------------------+",
        "        |                           |",
        "        v                           v",
        "  node_dq_reconcile         node_weather_multi_corridor",
        "  (Feature 3)               (Feature 5a)",
        "  - DQ-01 to DQ-04          - 9 waypoints, 2 corridors",
        "  - Item Master lookup       - Max-by-day corridor risk",
        "  - PoP trend analysis       - 48h risk and travel buffer",
        "        |                           |",
        "        +---------------------------+",
        "                     |",
        "                     v",
        "          node_resource_allocator",
        "          (Feature 5b)",
        "          - Load availability CSV",
        "          - Compute demand per corridor per day",
        "          - Penalty-minimising greedy allocation",
        "                     |",
        "                     v",
        "             node_planner  <----- retry",
        "                     |",
        "                     v",
        "               node_audit",
        "                     |",
        "          +----------+-------------------+",
        "          |          |                   |",
        "          v          v                   v",
        "       approved   escalate         human_checkpoint",
        "          |         (risk=3)             |",
        "          +----------+-------------------+",
        "                     |",
        "                     v",
        "               node_report --> node_email",
    ])

    heading2(doc, "3.2  Agent Design and Specialized Prompts")
    body(doc, "Five LLM agents operate in the pipeline, each with a dedicated prompt template and a specific responsibility boundary. A critical design choice is that no agent is asked to compute numeric values. All quantitative work (DQ counts, truck requirements, resource allocation penalties) is performed by deterministic Python code. Agents receive those computed numbers as inputs and are responsible only for interpreting, narrating, or auditing them.")

    agent_tbl = make_header_table(doc, ["Agent", "Prompt", "Input", "Output"])
    agent_data = [
        ("ContextAgent",         "PDF_CONTEXT_PROMPT",       "RAG-retrieved policy snippets",                         "Structured business context (SLAs, constraints, thresholds)"),
        ("TrendOpsAgent",        "TREND_OPS_PROMPT",         "DQ report, corridor KPIs, PoP deltas",                  "5-section markdown operations brief"),
        ("ResourcePlannerAgent", "RESOURCE_PLANNER_PROMPT",  "Availability, demand, allocation result, weather",      "Plain-English resource narrative for the planner"),
        ("PlannerAgent",         "PLANNER_PROMPT",           "Context, ops insights, weather risk, prior violations", "Strict JSON: buffer %, escalation flag, cited rules, plan text"),
        ("AuditAgent",           "AUDIT_PROMPT",             "Context, dispatch plan, weather risk, prior violations", "Strict JSON: passed boolean, violations list, severity"),
        ("ReportAgent",          "REPORT_PROMPT",            "All prior outputs, audit trail",                        "Full HTML dispatch report for leadership"),
    ]
    for i, row in enumerate(agent_data):
        add_data_row(agent_tbl, list(row), shade_even=True, row_idx=i)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    body(doc, "Two prompts have explicit anti-hallucination constraints. TREND_OPS_PROMPT instructs the agent: 'DO NOT invent numbers. Every claim must reference a value present in the provided inputs.' RESOURCE_PLANNER_PROMPT carries the same constraint. This makes agent output auditable because every number in the narrative can be traced back to a deterministic calculation.")

    callout(doc, "Two-Layer Audit Design: The AuditAgent (LLM) checks six qualitative rules about plan content. The deterministic audit layer (Python code) then independently checks three numeric facts: the buffer percentage exactly matches the Playbook section 5.2 mapping for the current weather risk score, escalation_required is true when and only when risk score equals 3, and cited_rules is non-empty. The plan passes only if both layers agree. This makes self-correction verifiable by code, not by asking the LLM to judge its own output.")

    heading2(doc, "3.3  Data Quality and Reconciliation Engine")
    body(doc, "The item master reconciliation engine is implemented in src/tools/dq_tools.py as pure, deterministic Python. It contains three reference tables hard-coded as module constants: Appendix A.1 (11 canonical items with item IDs, medicine types, temperature requirements, and product classes), Appendix A.2 (7 name alias mappings covering space variants, brand names, abbreviations, and spelling variants), and Appendix A.3 (4 legacy or deprecated item ID mappings).")
    arch_box(doc, [
        "Decision Tree per Row (D1 to D8 from Playbook Appendix A.6):",
        "",
        "  Row arrives",
        "    |",
        "    +-- unique_item_id is null or empty?         --> DQ-01: excluded",
        "    |",
        "    +-- (item_id, item_name) exact hit in A.1?   --> EXACT_MATCH (valid)",
        "    |",
        "    +-- item_name matches A.2 alias table?",
        "    |     |",
        "    |     +-- alias canonical agrees with row item_id?   --> ALIAS_MATCH (valid)",
        "    |     |",
        "    |     +-- alias canonical disagrees but item_id ok?  --> DQ-03 (flagged)",
        "    |     |",
        "    |     +-- item_id unknown?                           --> ALIAS_MATCH (valid)",
        "    |",
        "    +-- item_id in A.3 legacy map?               --> LEGACY_ID_MAP (valid)",
        "    |",
        "    +-- item_id in master but name unclear?      --> DQ-03 (flagged)",
        "    |",
        "    +-- item_id not recognized anywhere?         --> DQ-02: excluded",
        "",
        "  Post-resolution:",
        "    +-- duplicate unique_item_id detected?       --> DQ-04: flagged",
    ])

    heading2(doc, "3.4  Trend Analysis and KPI Computation")
    body(doc, "The trend module (src/tools/trend_tools.py) computes four categories of metrics from the reconciled DataFrame. Corridor KPIs are computed separately for the planning window (Day 0 and Day 1), the full history window, and all rows combined. Period-over-Period trend splits the history window at its date midpoint into Period A (earlier half) and Period B (later half), producing absolute and percentage deltas per KPI per corridor. Item spike detection compares each item's planning-window unit count against its historical daily average, with a spike ratio above 1.5 flagged. Planning-window rows are strictly excluded from the baseline computation to avoid contaminating the historical average with current demand.")

    heading2(doc, "3.5  Resource Allocation")
    body(doc, "The resource allocator (src/tools/resource_tools.py) implements the penalty-minimising allocation required by Playbook section 13.2. For each planning day, the algorithm computes the constraining ratio as the minimum of the three resource availability ratios (standard trucks, temperature-controlled trucks, drivers). All corridors are scaled uniformly by this ratio, and unmet units are attributed to the highest-penalty demand category first.")

    pen_tbl = make_header_table(doc, ["Violation Type", "Penalty per Unit", "Priority Rank"])
    pen_data = [
        ("Tier 1 cold-chain SLA violation",       "180 points (100 + 80)", "1 (highest)"),
        ("Tier 1 room-temperature SLA violation",  "100 points",            "2"),
        ("Tier 2 cold-chain violation",            "120 points (40 + 80)",  "3"),
        ("Tier 2 room-temperature SLA violation",  "40 points",             "4"),
        ("Non-SLA delivery delay",                 "10 points",             "5 (lowest)"),
    ]
    for i, row in enumerate(pen_data):
        add_data_row(pen_tbl, list(row), shade_even=True, row_idx=i)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_page_break(doc)

    # ============================================================ SECTION 4
    heading1(doc, "4.  Results and Validation")

    heading2(doc, "4.1  Actual Pipeline Output")
    body(doc, "The following results are taken directly from a complete pipeline run against the 14-day multi-corridor shipment CSV (129 rows across two corridors). The planning window covers March 6 and March 7, 2026 (Day 0 and Day 1).")

    heading3(doc, "Planning Window KPIs by Corridor")
    kpi_tbl = make_header_table(doc, ["Corridor", "Total", "Valid", "Excluded", "Excl. Rate", "Cold Chain", "Cold %", "Tier 1", "Tier 1 %", "Std Trucks", "Temp Trucks", "Drivers"])
    kpi_data = [
        ("C1_I95_NJ_BOS", "18", "16", "2", "11.1%", "9", "56.3%", "12", "75.0%", "1", "1", "2"),
        ("C2_NJ_PHL",     "15", "14", "1", "6.7%",  "7", "50.0%", "11", "78.6%", "1", "1", "2"),
    ]
    ci = list(range(1, 12))
    for i, row in enumerate(kpi_data):
        add_data_row(kpi_tbl, list(row), center_cols=ci, shade_even=True, row_idx=i)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    heading3(doc, "Weather Risk Summary")
    wx_tbl = make_header_table(doc, ["Corridor", "Day 0 Risk", "Day 1 Risk", "48h Risk", "Travel Buffer", "Escalation"])
    wx_data = [
        ("C1_I95_NJ_BOS", "0", "0", "0", "0%", "No"),
        ("C2_NJ_PHL",     "0", "0", "0", "0%", "No"),
    ]
    for i, row in enumerate(wx_data):
        add_data_row(wx_tbl, list(row), center_cols=list(range(1, 6)), shade_even=True, row_idx=i)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    body(doc, "All nine waypoints across both corridors registered zero triggered conditions during this run. This maps to a zero-buffer requirement, which the PlannerAgent correctly identified and was confirmed by the deterministic audit check.")

    heading2(doc, "4.2  Data Quality Results")

    dq_stat_tbl = doc.add_table(rows=2, cols=4)
    dq_stat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    dqh = ["Total Rows", "Valid (96.1%)", "Excluded DQ-01", "DQ-02 / 03 / 04"]
    dqv = ["129", "124", "5", "0"]
    dqc = ["003366", "1A6E2E", "8B5E00", "1A6E2E"]
    for i, (h, v, c) in enumerate(zip(dqh, dqv, dqc)):
        hc = dq_stat_tbl.rows[0].cells[i]
        vc = dq_stat_tbl.rows[1].cells[i]
        _set_cell_bg(hc, c)
        _set_cell_bg(vc, c)
        hp = hc.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(v)
        hr.font.size = Pt(16); hr.bold = True; hr.font.color.rgb = WHITE
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.add_run(h)
        vr.font.size = Pt(8); vr.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    dq_tbl = make_header_table(doc, ["Rule", "Count", "Description", "Impact"])
    dq_rows = [
        ("DQ-01", "5", "Missing unique_item_id (3 in Day 0, 2 in History)", "Excluded from all dispatch calculations per Playbook section 12"),
        ("DQ-02", "0", "Item ID not in master or legacy table", "None"),
        ("DQ-03", "0", "Item name conflicts with item ID canonical", "None"),
        ("DQ-04", "0", "Duplicate unique_item_id", "None"),
    ]
    for i, row in enumerate(dq_rows):
        r = dq_tbl.add_row()
        for j, val in enumerate(row):
            p = r.cells[j].paragraphs[0]
            if j == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(9)
            if j == 1 and val != "0":
                run.font.color.rgb = RED_FLAG
                run.bold = True
            elif j == 1 and val == "0":
                run.font.color.rgb = GREEN_OK
        if i % 2 == 1:
            for j in range(4):
                _set_cell_bg(r.cells[j], "F5F8FF")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    heading3(doc, "Alias Correction Example (from actual output)")
    alias_tbl = make_header_table(doc, ["Raw item_id", "Name in CSV", "Resolved Canonical ID", "Canonical Name", "Match Type"])
    alias_data = [
        ("10021", "Remdesivir 100 mg",       "RMD-100", "Remdesivir 100mg",           "alias_match"),
        ("10040", "EpiPen Auto Injector",     "EPI-AI",  "Epinephrine Auto-Injector",  "alias_match"),
        ("10050", "Heparin Na",               "HEP-SOD", "Heparin Sodium",             "alias_match"),
        ("10060", "Morphine Sulphate",        "MOR-SUL", "Morphine Sulfate",           "alias_match"),
        ("10070", "Albuterol Inhaler 90mcg",  "ALB-INH", "Albuterol Inhaler",          "alias_match"),
    ]
    for i, row in enumerate(alias_data):
        add_data_row(alias_tbl, list(row), shade_even=True, row_idx=i)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    heading2(doc, "4.3  Item Spike Detection")
    body(doc, "The spike analysis compared each item's planning-window unit count against its average daily units across the 12-day history window. All items with a spike ratio at or above 1.5 are flagged. The following table is taken directly from the deep-dive appendix generated by this pipeline run.")
    spike_tbl = make_header_table(doc, ["Item", "Planning Window Units", "Historical Daily Avg", "Spike Ratio"])
    spike_data = [
        ("Albuterol Inhaler",                          "4", "1.67", "2.40x"),
        ("Epinephrine Auto-Injector",                  "4", "1.75", "2.29x"),
        ("Heparin Sodium",                             "3", "1.29", "2.33x"),
        ("Insulin Lispro",                             "4", "1.80", "2.22x"),
        ("Pembrolizumab",                              "4", "1.86", "2.15x"),
        ("Remdesivir 100mg",                           "3", "1.44", "2.08x"),
        ("Morphine Sulfate",                           "3", "1.50", "2.00x"),
        ("Remdesivir 200mg",                           "2", "1.00", "2.00x"),
        ("Experimental Oncology Drug (Clinical Trial)", "2", "0.00", "No history"),
    ]
    for i, row in enumerate(spike_data):
        r = spike_tbl.add_row()
        for j, val in enumerate(row):
            p = r.cells[j].paragraphs[0]
            if j > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(9)
            if j == 3:
                run.font.color.rgb = RED_FLAG
                run.bold = True
        if i % 2 == 1:
            for j in range(4):
                _set_cell_bg(r.cells[j], "F5F8FF")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    heading2(doc, "4.4  Audit Loop Validation")
    heading3(doc, "Live Run Log (Console Output)")
    arch_box(doc, [
        r"[RAGEval] Recall@5=1.0 grounded_accuracy=1.0",
        "[ContextAgent] Cache hit -- reusing business_context.",
        "[DQReconcile] total=129 valid=124 excluded=5 flagged=0",
        "               (DQ-01=5, DQ-02=0, DQ-03=0, DQ-04=0)",
        "[Timing] dq_reconcile took 7.04s",
        "[Weather] C1_I95_NJ_BOS: Day0=0 Day1=0 48h=0 buffer=0% (5 waypoints)",
        "[Weather] C2_NJ_PHL: Day0=0 Day1=0 48h=0 buffer=0% (4 waypoints)",
        "[Timing] weather_multi_corridor took 2.10s",
        "[Timing] resource_allocator took 4.30s",
        "[AuditAgent] Attempt 1/3: PASSED",
        "[Timing] audit took 1.86s",
        "[Timing] report took 22.40s",
        "[Output] Full report saved to report_output.html",
        "[Output] Deep-dive appendix saved to report_appendix.txt",
    ])

    heading3(doc, "Unit Test Run Output")
    body(doc, "The following is the full pytest output from a live run of the complete test suite (29 tests across all five test files). The dark terminal block below shows every test name, its PASSED status, and the progress percentage.")
    add_terminal(doc)

    heading3(doc, "Unit Test Coverage")
    ut_tbl = make_header_table(doc, ["Test File", "Tests", "What Is Verified"])
    ut_data = [
        ("test_audit_logic.py",    "6",  "Buffer percentage mapping, pass/fail logic, escalation at score 3, empty cited_rules violation, LLM failure propagation"),
        ("test_dq_tools.py",       "10", "Exact match, alias match, legacy ID map, D4 precedence over D5, DQ-01 through DQ-04, SLA tier back-fill, missing column error"),
        ("test_resource_tools.py", "8",  "CSV loading, capacity model math, cold-chain routing, DQ exclusion, zero-penalty at full capacity, shortage penalty, bottleneck identification"),
        ("test_trend_tools.py",    "4",  "Capacity model formula, corridor KPI computation, PoP trend availability, deep-dive table shapes"),
        ("test_smoke.py",          "1",  "Import and graph assembly without OPENAI_API_KEY"),
        ("Total",                  "29", "All passing, 0 failures"),
    ]
    for i, row in enumerate(ut_data):
        r = ut_tbl.add_row()
        for j, val in enumerate(row):
            p = r.cells[j].paragraphs[0]
            if j == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(9)
            if i == 5:
                run.bold = True
                run.font.color.rgb = GREEN_DARK
        if i % 2 == 1:
            for j in range(3):
                _set_cell_bg(r.cells[j], "F5F8FF")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    heading2(doc, "4.5  Performance Benchmarks")
    perf_tbl = make_header_table(doc, ["Node", "Cold Run", "Warm Run", "Note"])
    perf_data = [
        ("pdf_context",              "14.69s", "0.20s",  "73x faster on warm run; ContextAgent LLM skipped via fingerprint cache"),
        ("dq_reconcile",             "7.04s",  "9.04s",  "Dominated by TrendOpsAgent LLM; pure DQ reconciliation is under 100ms"),
        ("weather_multi_corridor",   "2.10s",  "2.10s",  "9 HTTP calls to Open-Meteo, running sequentially"),
        ("resource_allocator",       "4.30s",  "4.30s",  "Dominated by ResourcePlannerAgent LLM"),
        ("planner",                  "16.91s", "15.30s", "JSON extraction and structured output parsing included"),
        ("audit",                    "1.86s",  "0.93s",  "First-attempt pass; deterministic checks add negligible overhead"),
        ("report",                   "22.40s", "23.79s", "ReportAgent LLM; largest single node"),
        ("Total",                    "~63s",   "~50s",   ""),
    ]
    for i, row in enumerate(perf_data):
        r = perf_tbl.add_row()
        for j, val in enumerate(row):
            p = r.cells[j].paragraphs[0]
            if j in (1, 2):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(9)
            if i == 7:
                run.bold = True
        if i % 2 == 1:
            for j in range(4):
                _set_cell_bg(r.cells[j], "F5F8FF")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_page_break(doc)

    # ============================================================ SECTION 5
    heading1(doc, "5.  Limitations and Next Steps")

    heading2(doc, "Current Constraints")

    heading3(doc, "Static Item Master")
    body(doc, "The canonical item master, alias table, and legacy ID mapping are hard-coded as Python module constants. This was the right choice for a controlled academic project where the reference data is known and stable, but it introduces a maintenance burden in production. Adding a new medicine or an alias requires a code change and a redeployment. In a real operations environment, these tables would need to be maintained in a database or configuration file by the pharmacy or procurement team without engineering involvement.")

    heading3(doc, "Sequential Weather Fetching")
    body(doc, "The multi-corridor weather node currently fetches nine waypoints sequentially, one HTTP call at a time. The total fetch time is approximately 2 seconds, which is acceptable today. If the corridor catalog expands to four or five corridors with additional waypoints, this will scale linearly and could become a meaningful bottleneck. The fix is straightforward: replace the sequential loop with concurrent HTTP requests using Python's asyncio or a thread pool executor, which would compress all nine calls into roughly one round-trip time.")

    heading3(doc, "LLM Non-Determinism and Cost")
    body(doc, "Five of the eight pipeline nodes invoke GPT-4.1-mini. While the system is designed so that no agent computes numbers (only narrates or audits them), the LLM responses still vary slightly across runs. This means that two identical input datasets could theoretically produce different narrative text in the report, and in edge cases, an audit that passes on one run might fail on another if the LLM agent's tone shifts. The deterministic audit layer catches numeric violations reliably, but qualitative judgment calls by the AuditAgent remain probabilistic.")

    heading3(doc, "Single-Tenant, Single-Run Design")
    body(doc, "The current system is designed to run once per day as a standalone script. It has no web interface, no scheduling mechanism, and no concurrent user support. The ChromaDB vector store is a local persistent directory and would not handle concurrent writes from multiple pipeline instances. For a multi-user production deployment, the Chroma instance would need to be replaced with a hosted vector database service.")

    heading2(doc, "Scaling with Real-World Data")

    heading3(doc, "Expanding the Corridor Network")
    body(doc, "The system is architecturally ready for additional corridors. Adding a new corridor requires three changes: adding the corridor's waypoints to the _CORRIDORS dictionary in weather_tools.py, adding a row to the corridor catalog table in the Dispatch Playbook, and ensuring the new corridor ID appears correctly in the shipment CSV. The DQ reconciliation, KPI computation, and resource allocation code all operate dynamically over whatever corridors appear in the data.")

    heading3(doc, "Higher Shipment Volumes")
    body(doc, "The reconciliation engine is vectorized and would handle 10,000 or 100,000 row CSVs without code changes. The binding constraint at scale would be the LLM calls: the TrendOpsAgent receives the full corridor comparison and PoP trend as a dictionary, which grows larger as the number of corridors and history days increases. Summarising the analytics to key signals before sending to the LLM, rather than passing the full dictionary, would keep prompt sizes manageable.")

    heading3(doc, "Real-Time and Streaming Data")
    body(doc, "The current architecture pulls a static CSV file at run time. A production enhancement would stream incoming shipment events from a message queue (for example, Apache Kafka or AWS Kinesis) and run the DQ reconciliation node on each batch. The trend analysis node could be adapted to maintain a rolling aggregate rather than recomputing from scratch on each run. The LangGraph StateGraph is already designed around immutable state transitions and would support an event-driven execution model with minimal architectural changes.")

    heading3(doc, "ML-Based Anomaly Detection and Feedback Loop")
    body(doc, "The current spike detection uses a simple ratio comparison against the historical daily average. A time-series forecasting model such as Facebook Prophet or a simple ARIMA model trained on the rolling history would produce a more statistically grounded baseline. The most valuable production enhancement would be closing the loop between dispatch recommendations and actual outcomes. If the system could receive post-delivery data (actual transit times, SLA compliance results, any temperature excursions for cold-chain items), it could compare the audit-approved plan against reality and surface systematic gaps in the planner's assumptions over time.")

    # Footer paragraph
    doc.add_paragraph()
    fp = doc.add_paragraph("SeeWeeS Ops Reporting Agent  |  UCLA MSBA AI Agents Project Challenge 2026  |  Technical and Business Report  |  May 2026")
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in fp.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = GREY_TXT

    # ---------------------------------------------------------------- Save
    docx_path = os.path.join(os.path.dirname(__file__), "Technical_Business_Report.docx")
    doc.save(docx_path)
    print(f"[OK] DOCX saved: {docx_path}")
    return docx_path


# ---------------------------------------------------------------------------
# PDF conversion
# ---------------------------------------------------------------------------
def convert_to_pdf(docx_path: str):
    pdf_path = docx_path.replace(".docx", ".pdf")
    try:
        from docx2pdf import convert
        print("[..] Converting to PDF via Microsoft Word...")
        convert(docx_path, pdf_path)
        print(f"[OK] PDF saved: {pdf_path}")
    except Exception as e:
        print(f"[!!] PDF conversion failed: {e}")
        print("     Open Technical_Business_Report.docx in Word and File > Export > PDF.")


if __name__ == "__main__":
    docx_path = build("Technical_Business_Report.docx")
    convert_to_pdf(docx_path)
